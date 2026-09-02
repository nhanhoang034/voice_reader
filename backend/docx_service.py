import io
import re
from docx import Document
from text_utils import split_and_log_sentences

def extract_and_clean_docx(file_bytes: bytes) -> dict:
    doc = Document(io.BytesIO(file_bytes))
    all_lines = []

    for p in doc.paragraphs:
        txt = p.text.strip()
        if not txt:
            continue

        # 1. Bỏ qua hẳn đoạn chú thích hình ảnh (bắt đầu bằng "Hình:" hoặc "Figure:")
        if re.match(r'^(Hình|Figure)\s*:', txt, flags=re.IGNORECASE):
            continue

        # 2. Bỏ qua các dòng chỉ chứa URL
        if re.match(r'^https?://\S+$', txt):
            continue

        # 3. Bỏ qua các dòng footnote ở đáy văn bản (dạng "(1) https://..." hoặc "(5)Bộ Luật...")
        if re.match(r'^\(\d+\)\s*(https?://|Bộ\s+Luật)', txt, flags=re.IGNORECASE):
            continue

        all_lines.append(txt)

    # Đọc bảng (nếu có)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text.strip()
                if txt and txt not in all_lines:
                    all_lines.append(txt)

    full_text = " ".join(all_lines)
    sentences = split_and_log_sentences(full_text, page_num=1, label="WORD")

    return {"pages": 1, "items": sentences}